#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Extract readable PRD content directly from a Word (.docx) document.

This script is intentionally read-only: it lists sections or prints extracted
content to stdout. It does not generate Markdown files.
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("缺少依赖：请先运行 pip install python-docx")

from case_utils import build_source_path, configure_output_encoding, ensure_under, project_root


HEADING_STYLE_PREFIXES = ("Heading", "heading", "标题")
LIST_BULLET_PREFIXES = ("List Bullet", "List bullet", "列表项目符号")
LIST_NUMBER_PREFIXES = ("List Number", "List number", "列表编号")


@dataclass(frozen=True)
class DocBlock:
    kind: str
    text: str
    level: int = 0


def _style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def _heading_level(paragraph) -> int:
    style_name = _style_name(paragraph)
    for prefix in HEADING_STYLE_PREFIXES:
        if style_name.startswith(prefix):
            suffix = style_name[len(prefix):].strip()
            if suffix.isdigit():
                return min(int(suffix), 6)
            if not suffix:
                return 1
    return 0


def _is_list_bullet(paragraph) -> bool:
    style_name = _style_name(paragraph)
    return any(style_name.startswith(prefix) for prefix in LIST_BULLET_PREFIXES)


def _is_list_number(paragraph) -> bool:
    style_name = _style_name(paragraph)
    return any(style_name.startswith(prefix) for prefix in LIST_NUMBER_PREFIXES)


def _list_level(paragraph) -> int:
    numpr = paragraph._element.find(qn("w:numPr"))
    if numpr is None:
        return 0
    ilvl = numpr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    value = ilvl.get(qn("w:val"))
    return int(value) if value and value.isdigit() else 0


def _has_inline_image(paragraph) -> bool:
    xml = paragraph._element.xml
    return "w:drawing" in xml or "w:pict" in xml


def _get_image_rel_ids(paragraph) -> list[str]:
    """从段落元素中提取所有图片的 relationship ID。"""
    rel_ids: list[str] = []
    for blip in paragraph._element.iter(qn("a:blip")):
        embed = blip.get(qn("r:embed"))
        if embed:
            rel_ids.append(embed)
    return rel_ids


def _get_image_extension(image_part) -> str:
    """根据 image_part 的 content_type 获取文件扩展名。"""
    content_type = getattr(image_part, "content_type", "")
    type_map = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tif",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
    }
    return type_map.get(content_type, ".png")


def _paragraph_text(paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        run_xml = run._element.xml
        if "w:drawing" in run_xml or "w:pict" in run_xml:
            continue
        parts.append(run.text)
    return "".join(parts).strip()


def _table_lines(table) -> list[str]:
    lines: list[str] = []
    for row_index, row in enumerate(table.rows, start=1):
        cells = []
        for cell in row.cells:
            cell_text = " ".join(
                text
                for text in (_paragraph_text(paragraph) for paragraph in cell.paragraphs)
                if text
            )
            cells.append(cell_text)
        lines.append(f"表格行{row_index}：" + " | ".join(cells))
    return lines


def read_docx_blocks(docx_path: Path) -> list[DocBlock]:
    doc = Document(str(docx_path))
    blocks: list[DocBlock] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._element is element), None)
            if paragraph is None:
                continue
            text = _paragraph_text(paragraph)
            if not text:
                continue

            level = _heading_level(paragraph)
            if level:
                blocks.append(DocBlock(kind="heading", text=text, level=level))
                continue

            if _is_list_bullet(paragraph):
                indent = "  " * _list_level(paragraph)
                blocks.append(DocBlock(kind="paragraph", text=f"{indent}- {text}"))
                continue

            if _is_list_number(paragraph):
                indent = "  " * _list_level(paragraph)
                blocks.append(DocBlock(kind="paragraph", text=f"{indent}1. {text}"))
                continue

            blocks.append(DocBlock(kind="paragraph", text=text))

        elif tag == "tbl":
            table = next((t for t in doc.tables if t._element is element), None)
            if table is None:
                continue
            for line in _table_lines(table):
                blocks.append(DocBlock(kind="table", text=line))

    return blocks


def list_sections(blocks: list[DocBlock]) -> list[DocBlock]:
    return [block for block in blocks if block.kind == "heading"]


def _normalized_heading(text: str) -> str:
    return " ".join(text.strip().lower().split())


def extract_section(blocks: list[DocBlock], section: str) -> list[DocBlock]:
    normalized = _normalized_heading(section)
    headings = [
        (index, block)
        for index, block in enumerate(blocks)
        if block.kind == "heading"
    ]
    matches = [
        (index, block)
        for index, block in headings
        if _normalized_heading(block.text) == normalized
    ]

    if not matches:
        matches = [
            (index, block)
            for index, block in headings
            if normalized in _normalized_heading(block.text)
        ]

    if not matches:
        return []
    if len(matches) > 1:
        matched_titles = "、".join(block.text for _, block in matches[:5])
        raise ValueError(f"章节名称不唯一，请输入更完整标题。匹配到：{matched_titles}")

    start_index, start_block = matches[0]
    start_level = start_block.level
    end_index = len(blocks)
    for index in range(start_index + 1, len(blocks)):
        block = blocks[index]
        if block.kind == "heading" and block.level <= start_level:
            end_index = index
            break

    return blocks[start_index:end_index]


def render_blocks(blocks: list[DocBlock]) -> str:
    lines: list[str] = []
    previous_kind = ""
    for block in blocks:
        if block.kind == "heading":
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(f"{'  ' * (block.level - 1)}{block.text}")
            lines.append("")
        else:
            if previous_kind == "heading" and lines and lines[-1] != "":
                lines.append("")
            lines.append(block.text)
        previous_kind = block.kind

    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines) + ("\n" if lines else "")


def section_has_images(docx_path: Path, section: str) -> bool:
    """Return True when the section contains at least one extractable image."""
    doc = Document(str(docx_path))
    normalized_section = _normalized_heading(section)
    section_started = False
    section_level = 0

    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        level = _heading_level(paragraph)

        if not section_started:
            if level and _normalized_heading(text) == normalized_section:
                section_started = True
                section_level = level
            continue

        if level and level <= section_level:
            break

        if _has_inline_image(paragraph) and _get_image_rel_ids(paragraph):
            return True

    return False


def extract_section_images(
    docx_path: Path, section: str, output_dir: Path
) -> list[Path]:
    """从指定章节中提取所有嵌入图片，保存到 output_dir 目录。

    返回保存的图片路径列表（不含已存在的图片）。
    """
    doc = Document(str(docx_path))
    normalized_section = _normalized_heading(section)
    saved_paths: list[Path] = []
    image_count = 0

    section_started = False
    section_level = 0

    for paragraph in doc.paragraphs:
        text = _paragraph_text(paragraph)
        level = _heading_level(paragraph)

        if not section_started:
            if level and _normalized_heading(text) == normalized_section:
                section_started = True
                section_level = level
            continue

        if level and level <= section_level:
            break

        if not _has_inline_image(paragraph):
            continue

        for rel_id in _get_image_rel_ids(paragraph):
            image_part = doc.part.related_parts.get(rel_id)
            if image_part is None or not hasattr(image_part, "blob"):
                continue
            image_count += 1
            ext = _get_image_extension(image_part)
            image_name = f"page_{image_count}{ext}"
            image_path = output_dir / image_name
            if image_path.exists():
                continue
            image_path.write_bytes(image_part.blob)
            saved_paths.append(image_path)

    return saved_paths


def discover_docx_files(root: Path) -> tuple[list[Path], list[Path]]:
    raw_docs = root / "inputs" / "requirements" / "raw_docs"
    if not raw_docs.exists():
        return [], []
    docx_files = sorted(raw_docs.glob("*.docx"))
    usable = [path for path in docx_files if not path.name.startswith("~$")]
    lock_files = [path for path in docx_files if path.name.startswith("~$")]
    return usable, lock_files


def print_docx_candidates(root: Path) -> None:
    usable, lock_files = discover_docx_files(root)
    if usable:
        print("可用 Word 文档：")
        for path in usable:
            print(f"  {path.relative_to(root)}")
    else:
        print("未找到可用 Word 文档。")

    if lock_files:
        print("\n已忽略 Word 临时锁文件：")
        for path in lock_files:
            print(f"  {path.relative_to(root)}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="从 Word (.docx) 直接列出章节或提取章节内容。"
    )
    parser.add_argument(
        "docx",
        nargs="?",
        help="源 Word 文件路径，例如 inputs/requirements/raw_docs/<文件名>.docx",
    )
    action_group = parser.add_mutually_exclusive_group()
    action_group.add_argument(
        "--section",
        default="",
        help="提取指定章节（优先精确匹配，唯一时允许标题包含匹配），不传则输出全文文本",
    )
    action_group.add_argument(
        "--list-sections",
        action="store_true",
        help="列出文档中所有章节标题",
    )
    action_group.add_argument(
        "--find-docs",
        action="store_true",
        help="列出 inputs/requirements/raw_docs/ 下可用的 .docx，并忽略 ~$ 临时锁文件",
    )
    parser.add_argument(
        "--print",
        action="store_true",
        help="兼容旧命令参数；当前脚本始终打印到标准输出",
    )
    parser.add_argument(
        "--extract-images",
        action="store_true",
        help="提取指定章节下的嵌入图片到 inputs/ui_design/<章节名>/ 目录，需配合 --section 使用",
    )
    return parser.parse_args()


def resolve_docx_path(raw_path: str) -> Path:
    root = project_root()
    docx_path = ensure_under(build_source_path(raw_path, root), root, "Word 文件")

    if not docx_path.exists():
        print_docx_candidates(root)
        raise FileNotFoundError(f"文件不存在：{docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx 格式，收到：{docx_path.suffix}")
    if docx_path.name.startswith("~$"):
        print_docx_candidates(root)
        raise ValueError(
            f"收到 Word 临时锁文件：{docx_path.name}，请改用不以 ~$ 开头的正式 .docx"
        )
    return docx_path


def main() -> int:
    configure_output_encoding()
    args = parse_args()
    root = project_root()

    if args.find_docs:
        print_docx_candidates(root)
        return 0

    if not args.docx:
        print("错误：请提供 Word 文件路径，或使用 --find-docs 查找可用文档。", file=sys.stderr)
        return 1

    try:
        docx_path = resolve_docx_path(args.docx)
        blocks = read_docx_blocks(docx_path)
    except (FileNotFoundError, ValueError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 1
    except Exception as exc:
        print(f"错误：读取 Word 文档失败：{exc}", file=sys.stderr)
        return 1

    if args.list_sections:
        sections = list_sections(blocks)
        if not sections:
            print("未找到任何标题章节。")
            return 0
        print(f"共 {len(sections)} 个章节：\n")
        for section in sections:
            print(f"{'  ' * (section.level - 1)}{'#' * section.level} {section.text}")
        return 0

    if args.extract_images:
        if not args.section:
            print("错误：--extract-images 需要配合 --section 使用", file=sys.stderr)
            return 1
        try:
            selected_blocks = extract_section(blocks, args.section)
        except ValueError as exc:
            print(f"错误：{exc}", file=sys.stderr)
            return 1
        if not selected_blocks:
            print(f"错误：未找到章节「{args.section}」，请检查标题名称。", file=sys.stderr)
            return 1

        root = project_root()
        section_dir_name = args.section.strip()
        output_dir = root / "inputs" / "ui_design" / section_dir_name

        if not section_has_images(docx_path, args.section):
            print(f"章节「{args.section}」未检测到嵌入图片。")
            return 0

        output_dir.mkdir(parents=True, exist_ok=True)
        saved = extract_section_images(docx_path, args.section, output_dir)
        if not saved:
            existing = list(output_dir.glob("*"))
            if existing:
                print(f"章节「{args.section}」的图片已存在，未提取新图片。")
            else:
                print(f"章节「{args.section}」未检测到嵌入图片。")
            return 0

        print(f"已提取 {len(saved)} 张图片到 {output_dir.relative_to(root)}：")
        for path in saved:
            print(f"  {path.relative_to(root)}")
        return 0

    try:
        selected_blocks = extract_section(blocks, args.section) if args.section else blocks
    except ValueError as exc:
        print(f"错误：{exc}", file=sys.stderr)
        return 1
    if args.section and not selected_blocks:
        print(f"错误：未找到章节「{args.section}」，请检查标题名称。", file=sys.stderr)
        return 1

    print(render_blocks(selected_blocks), end="")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
